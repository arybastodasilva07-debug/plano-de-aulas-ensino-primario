import streamlit as st
import os
import io
import random
import string
import urllib.parse
from streamlit_pdf_viewer import pdf_viewer
from openai import OpenAI
from docx import Document
from docx.shared import Inches

#------------ TEMPO DE LOGIN 1 HORA -----------------------------------------------

@st.cache_data(ttl=3600)  # Lembra a validação por 1 hora
def validar_acesso_cache(email, senha):
    try:
        if email in st.secrets["PASSWORDS"] and senha == str(st.secrets["PASSWORDS"][email]):
            return True
    except:
        pass
    return False

# --- CONFIGURAÇÕES GLOBAIS (AQUI!) ---
BASE_DIR = "biblioteca_permanente"
FOLDER_DOCS = "Central_Documentos"

# ---------------- CONFIGURAÇÃO ----------------
st.set_page_config(page_title="Plano de Aula - INIDE Angola", layout="wide")
st.title("🇦🇴 SISTEMA PROFISSIONAL DE ELABORAÇÃO DE PLANO DE AULA")
st.subheader("Ensino Primário (Iniciação à 6ª Classe)")

# Estilo para esconder menus nativos e limpar a interface
st.markdown("""
    <style>
    .stAppDeployButton {display:none;}
    footer {visibility: hidden;}
    [data-testid="stHeader"] {background: rgba(0,0,0,0);}
    </style>
    """, unsafe_allow_html=True)

# ---------------- 2. FUNÇÕES DE APOIO ----------------
def gerar_codigo():
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
    

# ======== NÃO MEXER ================= # NÃO MEXE # =========== NÃO MEXER ==============


# ----------------------------------------------------------------
# 3. FUNÇÕES DE PERSISTÊNCIA E SEGURANÇA (Coloque antes do Login)
# ----------------------------------------------------------------

@st.cache_data(ttl=3600)  # Lembra a validação por 1 hora para evitar quedas
def validar_acesso_persistente(email, senha):
    try:
        # Verifica nos Secrets a lista [PASSWORDS]
        if email in st.secrets["PASSWORDS"] and senha == str(st.secrets["PASSWORDS"][email]):
            return True
    except Exception:
        return False
    return False

# ----------------------------------------------------------------
# 4. SISTEMA DE LOGIN INTEGRADO
# ----------------------------------------------------------------

if "autenticado" not in st.session_state:
    st.session_state.autenticado = False

if not st.session_state.autenticado:
    st.title("🇦🇴 Portal Pedagógico - Angola")
    
    # Script invisível para manter a conexão ativa com o servidor
    st.markdown("""
        <script>
        setInterval(function() {
            window.parent.postMessage({type: 'streamlit:set_component_value', value: Date.now()}, '*');
        }, 30000); 
        </script>
        """, unsafe_allow_html=True)

    tab_login, tab_solicitar = st.tabs(["🔐 Entrar", "🔑 Solicitar Acesso / Gerar Senha"])

    with tab_login:
        st.subheader("Acesso ao Portal")
        email_in = st.text_input("E-mail Google", placeholder="exemplo@gmail.com").strip().lower()
        pass_in = st.text_input("Chave de Acesso", type="password", placeholder="Insira sua chave")
        
        if st.button("Entrar no Portal"):
            if validar_acesso_persistente(email_in, pass_in):
                st.session_state.autenticado = True
                st.session_state.user_email = email_in
                st.success("Acesso autorizado! Carregando portal...")
                st.rerun()
            else:
                st.error("Credenciais inválidas. Verifique o e-mail e a chave.")

    with tab_solicitar:
        st.subheader("Solicitar Nova Chave")
        st.write("1. Introduza o seu e-mail para gerar o código de acesso no campo abaixo.")
        st.write("2. Clique no botão verde (Enviar Pedido via WhatsAppa) para solicitar o acesso ao Administrador via WhatsApp.")
        st.write("3. Entre em contacto com o Administrador pelos terminais: 948298246/954458413.")

        email_novo = st.text_input("Introduza o seu e-mail para o cadastramento aqui", key="reg_email")
        
        if email_novo:
            if "@" in email_novo:
                # Função gerar_codigo() deve estar definida no topo do seu script
                cod_sugerido = gerar_codigo()
                
                # Configuração do WhatsApp
                texto_whatsapp = f"Olá Prof. António Basto! Sou o professor(a) {email_novo}. Gostaria de adquirir o acesso ao Portal de Planos de Aula. Código de Referência: {cod_sugerido}"
                texto_url = urllib.parse.quote(texto_whatsapp)
                seu_numero = "244954458413" 
                link_wa = f"https://wa.me/{seu_numero}?text={texto_url}"
                
                st.info(f"O seu código gerado é: **{cod_sugerido}**")
                
                st.markdown(f"""
                    <a href="{link_wa}" target="_blank" style="text-decoration: none;">
                        <div style="background-color: #25D366; color: white; text-align: center; padding: 12px; border-radius: 8px; font-weight: bold; font-size: 16px;">
                            📱 Enviar Pedido via WhatsApp
                        </div>
                    </a>
                """, unsafe_allow_html=True)
            else:
                st.warning("Por favor, introduza um e-mail válido.")

    # Interrompe a execução aqui para quem não está logado
    st.stop()


# ======== NÃO MEXER ================= # NÃO MEXE # =========== NÃO MEXER ==============

# ---------------- 4. CONFIGURAÇÃO PÓS-LOGIN ----------------

# Se chegou aqui, o login foi um sucesso
is_gerente = (st.session_state.user_email == st.secrets["ADMIN_EMAIL"].strip())

# --- 1. CONFIGURAÇÃO DE PASTAS E NOMES (COLOQUE NO TOPO) ---
BASE_DIR = "biblioteca_permanente"

# Esta é a lista que o seu erro diz estar em falta:
CLASSES = ["Iniciação", "1ª Classe", "2ª Classe", "3ª Classe", "4ª Classe", "5ª Classe", "6ª Classe", "Gerais"]

# Esta lista é usada para a estrutura de documentos (sem o "Gerais")
CLASSES_NOMES = CLASSES[:-1] 

# Estrutura da Central de Documentos
ESTRUTURA_DOCS = {
    "Calendário Escolar": None,
    "Programas do Ensino Primário": CLASSES_NOMES,
    "Dosificações": CLASSES_NOMES,
    "Decretos": None,
    "Constituição da República": None,
    "Lei de Bases": None,
    "Regulamento Escolar": None,
    "Currículo por Nível": None,
    "Estatuto da Carreira Docente": None,
    "Estatuto do Ministério": None,
    "Cadernos de Avaliação": None,
    "Decretos Presidenciais": None,
    "Didáctica Geral e Pedagogia": None,
    "Outros Documentos": None
}

# Criar a estrutura fisicamente
for pasta, subpastas in ESTRUTURA_DOCS.items():
    p_path = os.path.join(BASE_DIR, "Documentos_Centrais", pasta)
    if not os.path.exists(p_path): os.makedirs(p_path)
    if subpastas:
        for sub in subpastas:
            s_path = os.path.join(p_path, sub)
            if not os.path.exists(s_path): os.makedirs(s_path)

# ---------------- 5. INTERFACE PRINCIPAL ----------------
with st.sidebar:
    st.header("⚙️ MENU")
    st.write(f"Usuário: **{st.session_state.user_email}**")
    if is_gerente: st.success("Modo Gerente Ativo")
    if st.button("Sair do Sistema"):
        st.session_state.autenticado = False
        st.rerun()


# --- 4. DEFINIÇÃO DE PERMISSÕES ---
# Verifica se o e-mail logado é o do administrador definido nos Secrets
is_gerente = (st.session_state.user_email == st.secrets["ADMIN_EMAIL"].strip())

#------------------DEFINIÇÃO DE LEITUA DOS MANUAI------------------------------

import base64

def exibir_pdf(caminho_arquivo):
    with open(caminho_arquivo, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    # Cria um iframe para exibir o PDF
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="800" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

# --- 5. ESTRUTURA DINÂMICA DE ABAS ---

# Definimos as abas básicas que todos veem
titulos_abas = [
    "📝 GERADOR DE PLANOS", 
    "📂 CENTRAL DE DOCUMENTOS", 
    "📚 LIVROS POR CLASSE"
]

# Se for o gerente, adicionamos a quarta aba de controle
if is_gerente:
    titulos_abas.append("📂 GERENCIAR ARQUIVOS")

# Criamos as abas fisicamente no Streamlit
abas = st.tabs(titulos_abas)

# --- ABA 1: GERADOR DE PLANOS (Acesso Total) ---
with abas[0]:
    st.header("📝 Criar Novo Plano de Aula")
    st.info("Configure os detalhes na barra lateral a esquerda para gerar novo plano clicando na seta »")
   

# ------------- ABA 2: CENTRAL DE DOCUMENTOS (Com Ler e Baixar) ------------------------

with abas[1]:
    st.header("📂 Central de Documentos Oficiais")
    st.write("Selecione uma categoria para visualizar ou baixar os documentos.")

    # Espaço reservado para o visualizador de PDF (aparece no topo quando clica em Ler)
    container_leitura_docs = st.container()

    caminho_central = os.path.join(BASE_DIR, FOLDER_DOCS)

    for categoria, subcategorias in ESTRUTURA_DOCS.items():
        with st.expander(f"📁 {categoria.upper()}"):
            caminho_cat = os.path.join(caminho_central, categoria)
            os.makedirs(caminho_cat, exist_ok=True)
            
            # CASO A: Pastas com Subpastas (ex: Programas e Dosificações)
            if subcategorias:
                for sub in subcategorias:
                    st.markdown(f"**📍 {sub}**")
                    c_sub = os.path.join(caminho_cat, sub)
                    os.makedirs(c_sub, exist_ok=True)
                    arquivos = os.listdir(c_sub)
                    
                    if not arquivos:
                        st.caption("Nenhum arquivo disponível.")
                    
                    for f in arquivos:
                        caminho_f = os.path.join(c_sub, f)
                        col_n, col_v, col_d, col_x = st.columns([2.5, 0.8, 0.8, 0.5])
                        
                        col_n.write(f"📄 {f}")
                        
                        # Botão LER
                        if col_v.button("👁️", key=f"v_{categoria}_{sub}_{f}", help="Ler documento"):
                            with container_leitura_docs:
                                st.subheader(f"Visualizando: {f}")
                                if st.button("❌ Fechar Visualização", key=f"close_{f}"):
                                    st.rerun()
                                exibir_pdf(caminho_f)
                                st.divider()

                        # Botão BAIXAR
                        with open(caminho_f, "rb") as db:
                            col_d.download_button("📥", db.read(), file_name=f, key=f"d_{categoria}_{sub}_{f}")
                        
                        # Botão APAGAR (Admin)
                        if is_gerente:
                            if col_x.button("🗑️", key=f"del_{categoria}_{sub}_{f}"):
                                os.remove(caminho_f)
                                st.rerun()
                    st.divider()
            
            # CASO B: Pastas Simples (Calendário, Decretos, etc.)
            else:
                arquivos = os.listdir(caminho_cat)
                if not arquivos:
                    st.caption("Nenhum arquivo disponível.")
                
                for f in arquivos:
                    caminho_f = os.path.join(caminho_cat, f)
                    col_n, col_v, col_d, col_x = st.columns([2.5, 0.8, 0.8, 0.5])
                    
                    col_n.write(f"📄 {f}")
                    
                    # Botão LER
                    if col_v.button("👁️", key=f"v_{categoria}_{f}", help="Ler documento"):
                        with container_leitura_docs:
                            st.subheader(f"Visualizando: {f}")
                            if st.button("❌ Fechar Visualização", key=f"close_s_{f}"):
                                st.rerun()
                            exibir_pdf(caminho_f)
                            st.divider()

                    # Botão BAIXAR
                    with open(caminho_f, "rb") as db:
                        col_d.download_button("📥", db.read(), file_name=f, key=f"d_{categoria}_{f}")
                    
                    # Botão APAGAR (Admin)
                    if is_gerente:
                        if col_x.button("🗑️", key=f"del_{categoria}_{f}"):
                            os.remove(caminho_f)
                            st.rerun()

# --- ABA 3: LIVROS POR CLASSE (Com Ler e Baixar) ---
with abas[2]:
    st.header("📚 Livros por Classe")
    
    # Criamos um espaço vazio no topo para exibir o livro quando "Ler" for clicado
    container_leitura = st.container()

    for classe in CLASSES_NOMES:
        with st.expander(f"📁 {classe.upper()}"):
            caminho_classe = os.path.join(BASE_DIR, classe)
            os.makedirs(caminho_classe, exist_ok=True)
            arquivos = os.listdir(caminho_classe)
            
            if not arquivos:
                st.caption("Nenhum manual disponível.")
            
            for arq in arquivos:
                caminho_livro = os.path.join(caminho_classe, arq)
                
                # Colunas: Nome | Ler | Baixar | Apagar
                col_txt, col_ver, col_down, col_del = st.columns([2.5, 0.8, 0.8, 0.5])
                
                col_txt.write(f"📖 {arq}")
                
                # Botão LER
                if col_ver.button("👁️ Ler", key=f"ver_{classe}_{arq}"):
                    with container_leitura:
                        st.subheader(f"Visualizando: {arq}")
                        if st.button("❌ Fechar Leitura"):
                            st.rerun()
                        exibir_pdf(caminho_livro)
                        st.divider()

                # Botão BAIXAR
                with open(caminho_livro, "rb") as f:
                    col_down.download_button("📥", f.read(), file_name=arq, key=f"d_{classe}_{arq}")
                
                # Botão APAGAR (Admin)
                if is_gerente:
                    if col_del.button("🗑️", key=f"del_{classe}_{arq}"):
                        os.remove(caminho_livro)
                        st.rerun()

# --- ABA 4: GERENCIAR ARQUIVOS (EXCLUSIVA DO GERENTE) ---
if is_gerente:
    with abas[3]:
        st.header("📂 Painel de Gestão de Conteúdo")
        st.info("Utilize esta área para carregar novos documentos para o portal.")

        # 1. Seleção do Destino Principal
        # Criamos uma lista com as categorias da Central de Documentos + as pastas dos Livros
        categorias_principais = list(ESTRUTURA_DOCS.keys()) + CLASSES_NOMES
        destino_cat = st.selectbox("1. Selecione a Categoria/Pasta Principal:", categorias_principais)

        # 2. Seleção de Subpasta (Se existir)
        # Verifica se a categoria selecionada tem subpastas (ex: Programas, Dosificações)
        subpasta_alvo = ""
        if destino_cat in ESTRUTURA_DOCS and ESTRUTURA_DOCS[destino_cat] is not None:
            subpasta_alvo = st.selectbox(f"2. Selecione a Classe para '{destino_cat}':", ESTRUTURA_DOCS[destino_cat])
        
        # 3. Upload do Ficheiro
        arquivo_upload = st.file_uploader("3. Selecione o ficheiro PDF:", type="pdf")

        if st.button("🚀 Confirmar e Salvar no Portal"):
            if arquivo_upload is not None:
                if destino_cat in CLASSES:
                    # Vai para a pasta de Livros (raiz da biblioteca)
                    diretorio_destino = os.path.join(BASE_DIR, destino_cat)
                else:
                    # Vai para a Central de Documentos usando a constante FOLDER_DOCS
                    if subpasta_alvo:
                        diretorio_destino = os.path.join(BASE_DIR, FOLDER_DOCS, destino_cat, subpasta_alvo)
                    else:
                        diretorio_destino = os.path.join(BASE_DIR, FOLDER_DOCS, destino_cat)

                # CRIAÇÃO FÍSICA E GRAVAÇÃO
                os.makedirs(diretorio_destino, exist_ok=True)
                caminho_final = os.path.join(diretorio_destino, arquivo_upload.name)
                
                with open(caminho_final, "wb") as f:
                    f.write(arquivo_upload.getbuffer())
                
                st.cache_data.clear() # Limpa o cache para forçar a leitura do novo ficheiro
                st.success(f"✅ Ficheiro disponível em: {destino_cat}")
                st.rerun()

        st.divider()
        st.subheader("📊 Resumo do Servidor")
        # Pequena estatística para o gerente
        total_arquivos = sum([len(files) for r, d, files in os.walk(BASE_DIR)])
        st.write(f"O portal contém atualmente **{total_arquivos}** documentos carregados.")


# --- SISTEMA DE ARMAZENAMENTO PERMANENTE ---

# Criar a pasta base se não existir
BASE_DIR = "biblioteca_permanente"
if not os.path.exists(BASE_DIR):
    os.makedirs(BASE_DIR)

CLASSES = ["Iniciação", "1ª Classe", "2ª Classe", "3ª Classe", "4ª Classe", "5ª Classe", "6ª Classe"]
for c in CLASSES:
    c_path = os.path.join(BASE_DIR, c)
    if not os.path.exists(c_path):
        os.makedirs(c_path)

# Pasta para documentos gerais
DOCS_GERAIS_DIR = os.path.join(BASE_DIR, "Gerais")
if not os.path.exists(DOCS_GERAIS_DIR):
    os.makedirs(DOCS_GERAIS_DIR)

# --- ESTADO DA SESSÃO PARA O VISUALIZADOR ---
if 'pdf_ativo' not in st.session_state:
    st.session_state.pdf_ativo = None
if 'nome_ativo' not in st.session_state:
    st.session_state.nome_ativo = None


# ======== NÃO MEXER ================= # NÃO MEXE # =========== NÃO MEXER ==============  CIMA


                                                #BAIXO BAIXO
                                                
#----------------NÃO MEXER---------NÃO MEXER--------NÃO MEXER---------NÃO MEXER----------NÃO MEXER----- BAIXO BAIXOOOOO

# --- BARRA LATERAL ---
with st.sidebar:
    st.title("⚙️ Dados da Aula")
    # Coloque aqui o resto das suas seleções...

# ---------------- OPENAI CLIENT ----------------
api_key = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
if not api_key:
    st.error("❌ API Key não configurada. Vá em Settings → Secrets e adicione OPENAI_API_KEY.")
    st.stop()

client = OpenAI(api_key=api_key)

# ---------------- CURRÍCULO COMPLETO ----------------
curriculo = {
    "Iniciação": {
        "Língua Portuguesa": {
            "TEMA 2 – A MINHA FAMÍLIA E EU": {
                "Estudo das Vogais": ["Letra I", "Letra O", "Letra U", "Letra E", "Letra A"],
                "Sons e Ditongos": ["Vogais nasais", "Ditongos orais", "Ditongos nasais"]
            },
            "TEMA 3 – EU VOU À ESCOLA": {
                "Consoantes Iniciais": ["Estudo da letra P", "Estudo da letra B", "Estudo da letra M"],
                "Consoantes Dentais": ["Estudo da letra T", "Estudo da letra D"]
            }
        },
        "Estudo do Meio": {
            "TEMA 1 - A DESCOBERTA DE SI PRÓPRIO": {
                "O Meu Corpo": ["Identificação pessoal", "Partes do corpo", "Órgãos dos sentidos"],
                "Higiene e Saúde": ["Higiene corporal", "Higiene alimentar", "Vacinas"]
            },
            "TEMA 5 - ALIMENTAÇÃO": {
                "Importância": ["Necessidade de alimentação", "Alimentação rica e variada"],
                "Origem e Cuidados": ["Fonte dos alimentos", "Cuidados a ter com os alimentos"]
            }
        },
        "Matemática": {
            "TEMA 2 – NÚMEROS E OPERAÇÕES": {
                "Números Naturais": ["Números de 1 a 10", "Números de 11 a 20", "A Dezena"],
                "Operações": ["Adição até 9", "Subtracção até 9", "Algoritmo vertical"]
            }
        }
    },
    # ----------------- 1ª Classe -----------------
    "1ª Classe": {
        "Língua Portuguesa": {
            "TEMA 1 - QUEM SOU EU?": [
                "Eu sou",
                "Eu chamo-me",
                "Identificar pessoas",
                "Expressões de delicadeza",
                "Grafismos",
                "História: O Sapo e o Ovo"
        ],
            "TEMA 2 – A MINHA FAMÍLIA E EU": [
                "A família da Ana",
                "O alfabeto",
                "Estudo da letra I",
                "Estudo da letra O",
                "Estudo da letra U",
                "Estudo da letra E",
                "Estudo da letra A",
                "Vogais nasais",
                "Ditongos orais",
                "Ditongos nasais"
        ],
            "TEMA 3 – EU VOU À ESCOLA": [
                "Estudo da letra P",
                "Estudo da letra B",
                "Estudo da letra M",
                "Estudo da letra T",
                "Estudo da letra D",
                "Estudo da letra Ç",
                "Estudo da letra S",
                "Estudo da letra Z"
        ],
            "TEMA 4 – O MEU CORPO E EU": [
                "Estudo da letra N",
                "Estudo da letra L",
                "Estudo da letra X",
                "Estudo da letra F",
                "Higiene corporal",
                "Estudo da letra V",
                "Estudo da letra G",
                "Estudo da letra J",
                "Estudo da letra C",
                "Estudo da letra Q",
                "Estudo da letra R"
            ],
            "TEMA 5 – OS ANIMAIS QUE EU CONHEÇO": [
                "Estudo da letra H",
                "Estudo da letra Y",
                "Estudo da letra W",
                "Som NH",
                "Som LH",
                "Som CH",
                "Som AR",
                "Som OL",
                "Som AM",
                "Som UM"
            ]
    },

        "Estudo do Meio": {
            "TEMA 1 - A DESCOBERTA DE SI PRÓPRIO": [
                "A minha identificação",
                "O meu corpo",
                "Órgãos dos sentidos",
                "Saúde do corpo",
                "Higiene do corpo",
                "Higiene alimentar",
                "Posturas correctas, repouso e sono",
                "Vacinas"
            ],
            "TEMA 2 - A FAMÍLIA": [
                "Membros da família",
                "Convivência com outras pessoas"
            ],
            "TEMA 3 - A HABITAÇÃO": [
                "Habitação",
                "Localização da habitação",
                "Compartimentos da habitação",
                "Habitação como lugar de convivência",
                "Cuidados com a casa"
            ],
            "TEMA 4 - A ESCOLA": [
                "A minha escola",
                "Localização da escola",
                "Compartimentos da escola",
                "A sala de aulas",
                "Cuidados a ter com a escola",
                "Cuidados a te com material escolar",
                "Comunidade escolar"
            ],
            "TEMA 5 - ALIMENTAÇÃO": [
                "Necessidade de alimentação",
                "Alimentação rica e variada",
                "O que devemos comer",
                "Alimentos presentes na nossa dieta",
                "Fonte dos alimentos",
                "Cuidados a ter com os alimentos"
            ],
            "TEMA 6 - O VESTUÁRIO": [
                "Tipos de vestuário",
                "Materiais têxteis",
                "Higiene do vestuário"
            ],
            "TEMA 7 - A SEGURANÇA": [
                "Itinerários",
                "Sinais de trânsito",
                "Regras de trânsito",
                "Cuidados com líquidos inflamáveis e objectos perigosos"
            ],
            "TEMA 8 – AS PLANTAS": [
                "Estrutura das plantas",
                "Importância das plantas",
                "Cuidados a ter com as plantas"
            ],
            "TEMA 9 – OS ANIMAIS": [
                "Animais domésticos e selvagens",
                "Estrutura dos animais",
                "Importância dos animais"
            ]
        },
    
        "Matemática": {
            "TEMA 1 – GEOMETRIA": [
                "Relações espaciais(À frente/atrás/entre; Em cima/em baixo; dentro/fora; Interios/Exterior; Alto/baixo; Direita/esquerda/à direita/à esquerda)",
                "Sólidos geométricos",
                "Figuras geométricas planas",
                "Linhas rectas, curvas e quebradas",
                "Linhas abertas e fechadas"
            ],
            "TEMA 2 – NÚMEROS, CONJUNTOS E OPERAÇÕES": [
                "Estudo dos números naturais até 10 (leitura e escrita)",
                "Adição de números naturais até 9",
                "Subtracção de números naturais até 9",
                "Comparação de números naturais até 10",
                "Leitura, escrita e adição dos números até 20: dezena e unidade",
                "Comparação e ordenação dos números naturais até 20"
                "Adição e subtracção dos números na forma do algoritmo vertical"
                "Conjuntos: Tantos como (quanto) | mais do que | menos do que",
                "Leitura e escrita dos números de 21 até 50",
                "Adição e subtracção de números naturais até 50",
                "Contar e escrever de 2 em 2, de 5 em 5 e de 10 em 10",
                "Leitura e escrita dos números de 51 até 100",
                "Adição e subtracção de números até 100",
                "Comparação e ordenação dos números até 100",
                "Composição e decomposição dos números em parcelas",
                "Multiplicação e divisão dos números naturais por 2, 3 e 4"
            ],
            "TEMA 3 – GRANDEZAS E MEDIDAS: Conservação, comparação e ordenação de grandezas": [
                "Comprimento (Comparação de grandezas – comprimento)",
                "Massa",
                "Capacidade",
                "Relações temporais: Hoje, ontem, amanhã, agora, antes, depois/ Muito tempo, pouco tempo, ao mesmo tempo",
                "Dias da semana",
                "Dinheiro e moeda angolana"
                "Valores faciais da moeda angolana até Kz 100.00",
            ]
        },
    
        "Educação Manual e Plástica": {
            "TEMA 1 – O DESENHO": [
                "Risco (desenhos com riscos)",
                "Traço e tracejo (desenhos com traços)",
                "Pontos (desenhos com pontos)",
                "Rasgagem e colagem (Rasgar e colar um triângulo e um quadrado,um círculo e uma cobra, uma estrela e uma casa, uma laranja e uma banana e várias figuras simples)"
            ],
            "TEMA 2 – A PINTURA": [
                "Impressão e estampagem (Estampar com frutos e folhas, Carimbar com plasticina, Imprimir no papel, Estampar com recortes)",
                "Digitintas (digitintas com as mãos, com os dedos e desenhar com os dedos)",
                "Recorte e colagem (Recortar as figuras geométricas simples utilizando a tesoura, Recortar uma maçã e um menino)"
            ],
            "TEMA 3 – A MODELAGEM": [
                "Preparar e amassar barro",
                "Separar e enrolar",
                "Modelagem com barro e plasticina"
            ],
            "TEMA 4 – AS CONSTRUÇÕES": [
                "Dobragem (Dobrar uma folha de papel em quatro partes iguais, Fazer um avião, fazer um chapéu e fazer um barco)",
                "Recorte e embrulho (Forrar um livro ou um caderno, embrulhar uma caixa)",
                "Composição e colagem (Fazer um quadro com folhas, Fazer uma flor, Compor uma árvore, Compor um animal (borboleta), Compor um animal (porco))",
                "Construção com materiais diversos (Construir uma casa com amorfo e plasticina, Construir um boneco)"
            ]
        },
    
        "Educação Musical": {
            "TEMA 1 – A VOZ": [
                "Pequenas canções",
                "Sons naturais e artificiais",
                "Canções populares locais e escolares"
            ],
            "TEMA 2 – O CORPO": [
                "Percussão corporal com canções e gestos",
                "Reprodução de batimentos simples e complexos"
            ],
            "TEMA 3 – INICIAÇÃO À TEORIA DA MÚSICA": [
                "Pauta musical",
                "Primeiro espaço suplementar inferior e primeira linha da pauta musical: notas Ré e Mi"
            ],
            "TEMA 4 – OS INSTRUMENTOS MUSICAIS": [
                "Instrumentos de percussão",
                "Audição dos sons produzidos por instrumentos de percussão",
                "Jogos musicais com instrumentos de percussão"
            ],
            "TEMA 5 – INICIAÇÃO À TEORIA DA MÚSICA II": [
                "Primeiro espaço e segunda linha da pauta musical: notas Fá e Sol"
            ],    
            "TEMA 6 – OS INSTRUMENTOS MUSICAIS": [
                "Audição dos sons produzidos por instrumentos de percussão"
            ]
        },
    
        "Educação Física": {
            "TEMA 1 – Ginástica básica": [
                "Deslocamento",
                "Salto",
                "Lançamento",
                "Trepar",
                "Equilíbrio"
            ],
            "TEMA 2 – Ginástica rítmica": [
                "Ritmo"
            ],
            "TEMA 3 – Atletismo": [
                "Corrida de velocidade",
                "Corrida de resistência"
            ],
            "TEMA 4 – Jogos": [
                "Jogos com bola",
                "Jogos de correr",
                "Jogos de saltar",
                "Jogos sensoriais"]
        }
    },
    # ----------------- 2ª Classe -----------------
    "2ª Classe": {
        "Língua Portuguesa": {
            "A Minha Escola": [
                "O encontro",
                "Pelo caminho",
                "A queda da Vera",
                "A Ana e o Paulo",
                "A escola limpa",
                "A casa do Guedes",
                "Um passeio",
                "As férias",
                "A sala de aulas",
                "Qual é a coisa qual é ela"
            ],
            "A Minha Família": [
                "A minha família",
                "O girassol",
                "A lavra",
                "O município de Cacuaco (I)",
                "O município de Cacuaco (II)",
                "O município de Cacuaco (III)",
                "Tenho fome",
                "O trabalho é importante",
                "Vamos cantar",
                "A avó da Ana",
                "Vamos recitar"
            ],
            "O Mundo dos Animais": [
                "Animais domésticos",
                "Animais selvagens",
                "Domésticos e selvagens",
                "A utilidade dos animais",
                "Tu falas",
                "O sonho da Glória",
                "As férias na aldeia (I)",
                "As férias na aldeia (II)",
                "As formigas",
                "Uma visita",
                "A trovoada"
            ],
            "A Minha Saúde": [
                "Os alimentos (I)",
                "Os alimentos (II)",
                "O Víctor está forte",
                "No campo",
                "Uma carta",
                "A chuva",
                "O mercado",
                "A saúde",
                "Higiene do corpo (I)",
                "Asseio",
                "Higiene do corpo (II)",
                "As vacinas",
                "Canções e provérbios"
            ]
        },
    
        "Estudo do Meio": {
            "A Descoberta de Ti Mesmo": [
                "O meu passado",
                "Meus gostos e preferências",
                "O meu corpo",
                "Cuido do meu corpo",
                "Segurança na via pública",
                "Doenças do meu meio"
            ],
            "A Família": [
                "A minha família",
                "Outros membros da família",
                "Relação de arentesco",
                "Necessidades da família"
            ],
            "A Habitação": [
                "Tipos de habitação",
                "Regras de higiene da habitação"
            ],
            "A Escola": [
                "Partes constituente da escola",
                "Cuidados a ter com o material escolar",
                "Higiene e saúde escolar"
            ],
            "A Alimentação": [
                "Fontes de alimentação",
                "Cuidados a ter com alimentação",
                "Higiene alimentar",
                "Conservação dos alimentos"
            ],
            "Necessidade do Vestuário": [
                "Importância do vestuário",
                "Tipos de vestuário",
                "Higiene do vestuário"
            ],
            "As Plantas": [
                "Plantas da localidade",
                "Plantas quanto ao tamanho: Árvores, arbustos e ervas",
                "Plantas espontâneas e semeadas",
                "Partes constituintes da planta",
                "Reprodução das plantas",
                "Importância das plantas",
                "Eu cuido do ambiente"
            ],
            "Os Animais": [
                "Animais da localidade",
                "Características externas de alguns animais",
                "Modo de vida de alguns animais",
                "Importância dos animais"
            ],
            "O Trabalho": [
                "O trabalho"
            ],
            "Transportes e Comunicações": [
                "Meios de transporte",
                "Meios de comunicação"
            ]
        },
    
        "Matemática": {
            "Geometria": [
                "Sólidos geométricos (Paralelepípedo, cubo, esfera, cilindro e cone",
                "Superfícies planas e curvas",
                "Figuras geométricas planas e os seus traçados:  Rectângulo, quadrado, triângulo e círculo",
                "Linhas directas, curvas e quebradas",
                "Itinerário e percurso"
            ],
            "Números e Operações": [
                "Estudo dos números naturais de 100 até 500 (leitura e escrita)",
                "Adição e subtracção de números naturais de 1 até 500",
                "Composição e decomposição de números naturais em parcelas",
                "Comparação e ordenação de números de 100 até 500",
                "Resolução de problemas envolvendo números naturais de 100 até 500",
                "Estudo dos números naturais de 500 até 1 000 (Leitura e escrita)",
                "Adição e subtracção de números naturais até 1000",
                "Composição e decomposição de números em parcelas",
                "Multiplicação por 2, 3, 4, 5 e por 10",
                "Divisão de números naturais por 2, 3, 4, 5 e 10"
                "Comparação e ordenação de números até 1 000",
                "Resolução de problemas envolvendo úmeros naturais de 500 até 1 000",
            ],
            "Grandezas e Medidas": [
                "Medidas não padronizadas: Comprimento (palmos, pé, passos e braços)",
                "Medidas não padronizadas: Capacidade (tanque, balde, lata e caneca)",
                "Medidas padronizadas (Metro, Grama e Capacidade(litro))",
                "Medidas de tempo (Os dias da semana, Os meses do ano)",
                "Medidas de tempo (Leitura da hora e do minuto no relógio)",
                "Resolução de problemas de medidas de tempo",
                "A moeda(Moeda angolana)",
                "A moeda (Valores faciais da moeda angolana até Kz 1000,00)",
                "Relação entre valores faciais da moeda",
                "Resolução de problemas que envolvam dinheiro até Kz 1 000,00,"
            ]
        },
    
        "Educação Manual e Plástica": {
            "Aprende a Desenhar": [
                "Desenho com o ponto",
                "Desenho com a linha"
            ],
            "Conhece as Cores": [
                "Desenho com as cores",
                "Aprende a pintar"
            ],
            "Cria Fora do Papel": [
                "Modelagem",
                "Reciclagem",
                "Técnicas mistas"
            ]
        },
    
        "Educação Musical": {
            "A Voz": [
                "Pequenas canções",
                "Sons naturais e artificiais",
                "Canções populares"
            ],
            "O Corpo": [
                "Percussão corporal",
                "Batimentos simples e complexos"
            ],
            "Teoria da Música": [
                "Pauta musical",
                "Notas Ré e Mi",
                "Notas Fá e Sol"
            ],
            "Instrumentos Musicais": [
                "Instrumentos de percussão",
                "Audição de sons",
                "Jogos musicais"
            ]
        },
    
        "Educação Física": {
            "Ginástica Básica": [
                "Deslocamento",
                "Salto",
                "Lançamento",
                "Trepar",
                "Equilíbrio"
            ],
            "Ginástica Rítmica": [
                "Ritmo"
            ],
            "Atletismo": [
                "Corrida de velocidade",
                "Corrida de resistência"
            ],
            "Jogos": [
                "Jogos com bola",
                "Jogos de correr",
                "Jogos de saltar",
                "Jogos sensoriais"]
         }
    },

    # ----------------- 3ª Classe -----------------
    "3ª Classe": {
        "Língua Portuguesa": {
            "Leitura e Compreensão": ["Textos curtos e pequenos contos", "Personagens e enredo"],
            "Escrita": ["Frases completas", "Dictados curtos"],
            "Gramática": ["Substantivos, adjetivos e pronomes", "Verbos no presente"]
        },
        "Matemática": {
            "Operações": ["Adição e subtração com dezenas", "Multiplicação inicial"],
            "Geometria e Medidas": ["Figuras planas e sólidas", "Medidas de comprimento e peso"]
        },
        "Estudo do Meio": {
            "Comunidade e Sociedade": ["Regras, profissões e funções", "Cooperação e cidadania"],
            "Meio Natural": ["Animais, plantas e ciclos da natureza"]
        },
        "Educação Visual e Plástica": {
            "Desenho e Pintura": ["Cores, linhas e formas mais complexas"],
            "Modelagem": ["Criação de formas e objetos simples"]
        },
        "Educação Física": {
            "Movimento": ["Coordenação, força e equilíbrio"],
            "Jogos e Brincadeiras": ["Trabalho em equipe e regras"]
        }
    },

    # ----------------- 4ª Classe -----------------
    "4ª Classe": {
        "Língua Portuguesa": {
            "Leitura": ["Textos narrativos e informativos", "Compreensão de enredo"],
            "Escrita": ["Redação curta", "Dictados intermediários"],
            "Gramática": ["Substantivos, adjetivos, pronomes e verbos"]
        },
        "Matemática": {
            "Números e Operações": ["Adição, subtração, multiplicação e divisão"],
            "Geometria": ["Figuras planas e sólidas", "Perímetro e área simples"]
        },
        "Estudo do Meio": {
            "Sociedade": ["Regras de convivência", "Cidadania"],
            "Meio Natural": ["Animais, plantas, água e solo"]
        },
        "Educação Visual e Plástica": {
            "Desenho e Pintura": ["Projetos de desenho detalhado", "Pintura de objetos e paisagens"],
            "Modelagem": ["Criação de figuras e esculturas simples"]
        },
        "Educação Física": {
            "Movimento": ["Força, coordenação e resistência"],
            "Jogos": ["Competição saudável e cooperação"]
        }
    },

    # ----------------- 5ª Classe -----------------
    "5ª Classe": {
        "Língua Portuguesa": {
            "Leitura": ["Textos literários e informativos", "Compreensão de personagens e enredo"],
            "Escrita": ["Redação média", "Dictado intermediário"],
            "Gramática": ["Classes de palavras, tempos verbais e pontuação"]
        },
        "Matemática": {
            "Números e Operações": ["Adição, subtração, multiplicação, divisão e frações"],
            "Geometria e Medidas": ["Figuras geométricas, perímetro, área e volume"]
        },
        "Estudo do Meio": {
            "Sociedade e Cidadania": ["Responsabilidade, regras e cidadania"],
            "Meio Natural": ["Ciclo da água, animais e plantas"]
        },
        "Educação Visual e Plástica": {
            "Desenho e Pintura": ["Projetos mais complexos", "Representação de paisagens"],
            "Modelagem": ["Esculturas simples e criativas"]
        },
        "Educação Física": {
            "Movimento": ["Coordenação avançada, resistência"],
            "Jogos": ["Cooperação, estratégia e regras"]
        }
    },

    # ----------------- 6ª Classe -----------------
    "6ª Classe": {
        "Língua Portuguesa": {
            "Leitura": ["Textos literários e científicos", "Compreensão crítica"],
            "Escrita": ["Redação avançada", "Resumo e ditado avançado"],
            "Gramática": ["Classes de palavras, tempos verbais e sintaxe"]
        },
        "Matemática": {
            "Números e Operações": ["Todas as operações, frações, decimais"],
            "Geometria e Medidas": ["Perímetro, área, volume e ângulos"]
        },
        "Estudo do Meio": {
            "Sociedade e Cidadania": ["Regras, ética e cidadania"],
            "Meio Natural": ["Ciclos da natureza, ecologia, sustentabilidade"]
        },
        "Educação Visual e Plástica": {
            "Desenho e Pintura": ["Projetos detalhados e criativos"],
            "Modelagem": ["Esculturas e criações tridimensionais"]
        },
        "Educação Física": {
            "Movimento": ["Coordenação, força, resistência e agilidade"],
            "Jogos": ["Estratégia, competição saudável e trabalho em grupo"]
        }
    }
}

# ---------------- SIDEBAR ----------------
with st.sidebar:
    st.header("🏫 Identificação")
    nome_escola = st.text_input("Nome da Escola")
    nome_professor = st.text_input("Nome do Professor")
    trimestre = st.selectbox("Trimestre", ["1º Trimestre", "2º Trimestre", "3º Trimestre"])
    logotipo = st.file_uploader("Logotipo da Escola (opcional)", type=["png","jpg","jpeg"])

    st.header("📚 Dados da Aula")
    classe = st.selectbox("Classe", list(curriculo.keys()))
    disciplinas = list(curriculo[classe].keys())
    disciplina = st.selectbox("Disciplina", disciplinas)
    temas = list(curriculo[classe][disciplina].keys())
    tema = st.selectbox("Tema", temas)
    subtemas = curriculo[classe][disciplina][tema]
    subtema = st.selectbox("Subtema", subtemas)
    sumários = curriculo[classe][disciplina][tema][subtema]
    sumário = st.selectbox("Sumário", sumários)
    aula_numero = st.number_input("Aula nº", min_value=1, step=1)
    tempo = "45 minutos"

    gerar = st.button("🧠 Gerar Plano de Aula")
    
# ----------------Coloque isso logo após a criação da sidebar---------------------
import time

# Script invisível para manter a conexão ativa
st.sidebar.markdown("""
    <script>
    setInterval(function() {
        window.parent.postMessage({type: 'streamlit:set_component_value', value: Date.now()}, '*');
    }, 30000); // Envia um sinal a cada 30 segundos
    </script>
    """, unsafe_allow_html=True)

# ---------------- FUNÇÃO PARA GERAR PLANO ----------------
def gerar_plano():
    prompt = f"""
Gere um plano de aula completo baseado no currículo oficial do INIDE (Angola).

Escola: {nome_escola}
Professor: {nome_professor}
Disciplina: {disciplina}
Classe: {classe}
Trimestre: {trimestre}
Tempo: {tempo}
Aula nº: {aula_numero}
Unidade temática: {tema}
Subtema: {subtema}
Sumário: {sumário}

Se o subtema for amplo, divida em mais de uma aula de 45 minutos.

Estrutura obrigatória:
1. Objetivo Geral
2. Objetivos da Aula
3. Conteúdo
4. Material Didáctico
5. Metodologia
6. Actividades Chave
7. Tipo de Avaliação
8. Procedimentos (O passo a passo):
    - Introdução (Acolhimento e apresentação ou motivação)
    - Desenvolvimento
    - Actividades ou exercícios
    - Consolidação
    - Avaliação
    - Tarefa para Casa

Linguagem formal pedagógica.
Contexto angolano.
"""
    try:
        resposta = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role":"system","content":"Especialista em planos de aula do ensino primário angolano."},
                {"role":"user","content":prompt}
            ],
            temperature=0.5,
            max_tokens=1500
        )
        return resposta.choices[0].message.content
    except Exception as e:
        st.error(f"❌ Não foi possível gerar o plano. Motivo: {e}")
        st.stop()

# ---------------- FUNÇÃO PARA GERAR WORD ----------------
def gerar_word(plano_texto):
    doc = Document()
    doc.add_heading("REPÚBLICA DE ANGOLA", level=1)
    doc.add_paragraph(f"Escola: {nome_escola}")
    doc.add_paragraph(f"Professor: {nome_professor}")
    doc.add_paragraph(f"Disciplina: {disciplina}")
    doc.add_paragraph(f"Classe: {classe}")
    doc.add_paragraph(f"Trimestre: {trimestre}")
    doc.add_paragraph(f"Aula nº: {aula_numero}")
    doc.add_paragraph(f"Tempo: {tempo}\n")
    doc.add_heading("PLANO DE AULA", level=2)
    doc.add_paragraph(plano_texto)
    if logotipo is not None:
        doc.add_picture(logotipo, width=Inches(1.5))
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------- EXECUÇÃO ----------------
if gerar:
    if not nome_escola or not nome_professor:
        st.warning("⚠️ Preencha Nome da Escola e Nome do Professor.")
    else:
        with st.spinner("Gerando plano profissional..."):
            plano = gerar_plano()
        st.success("✅ Plano gerado com sucesso!")
        st.markdown(plano)
        # Download TXT
        st.download_button("📥 Baixar em TXT", plano, "plano_de_aula.txt")
        # Download Word
        word_file = gerar_word(plano)
        st.download_button("📄 Baixar em Word (.docx)", word_file, "plano_de_aula.docx")







































































